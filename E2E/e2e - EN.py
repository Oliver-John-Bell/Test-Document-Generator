from openpyxl import load_workbook
import openpyxl
from openpyxl.utils import get_column_letter
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import pandas as pd

def regexfixv3(x):
    x = re.sub("_x000D_", "", x)
    def split_long_word(word):
        if '<<<' in word or '>>>' in word:
            parts = re.split(r'(<<<|>>>)', word)
            new_parts = []
            for part in parts:
                if len(part) > 19:
                    chunk_size = 19
                    new_parts.extend([part[i:i+chunk_size] for i in range(0, len(part), chunk_size)])
                else:
                    new_parts.append(part)
            return ' '.join(new_parts)
        elif len(word) > 19:
            return ' '.join([word[i:i+19] for i in range(0, len(word), 19)])
        else:
            return word

    processed_text = ''
    last_end = 0
    for match in re.finditer(r'\S+', x):
        start, end = match.span()
        processed_text += x[last_end:start]
        processed_text += split_long_word(match.group())
        last_end = end
    processed_text += x[last_end:]
    return processed_text

import re
import pandas as pd

class Step:
    def __init__(self, Scenario, TestID, Test_Name, Test_Description, Step_Name,
                 Step_Description, Expected_Results, Role, Workstream):
        self.Scenario = re.sub("_x000D_", "", str(Scenario))
        self.TestID = re.sub("_x000D_", "", str(TestID))
        self.TestName = re.sub("_x000D_", "", str(Test_Name))
        self.TestDescription = re.sub("_x000D_", "", str(Test_Description))
        self.StepName = self.regexfixv3(str(Step_Name))
        self.StepDescription = self.regexfixv3(str(Step_Description))
        self.ExpectedResults = self.regexfixv3(str(Expected_Results))
        self.Role = re.sub("_x000D_", "", str(Role))
        self.Workstream = re.sub("_x000D_", "", str(Workstream))
    
    def regexfixv3(self, text):
        return re.sub(r"_x000D_\n", "", text)

class Test:
    def __init__(self, Scenario, TestID, Test_Name, Test_Description):
        self.Scenario = Scenario  # Assign scenario directly to the Test
        self.TestID = TestID
        self.TestName = Test_Name
        self.TestDescription = Test_Description
        self.steps = []
    
    def add_step(self, step):
        self.steps.append(step)

def read_excel_to_tests(file_path):
    df = pd.read_excel(file_path)
    tests = []

    for _, row in df.iterrows():
        scenario_name = row['Scenario']
        TestID = row['TEST ID']
        Test_Name = row['TEST NAME']
        Test_Description = row['DESCRIPTION']
        Step_Name = row['Step Name']
        Step_Description = row['STEP DESCRIPTION']
        Expected_Results = row['Expected Results']
        Role = row['Role']
        Workstream = row['Workstream']

        # Find the test by its TestID and Scenario, or create a new one
        test = next((t for t in tests if t.TestID == TestID and t.Scenario == scenario_name), None)
        if test is None:
            test = Test(scenario_name, TestID, Test_Name, Test_Description)
            tests.append(test)
        
        # Create a step and add it to the test
        step = Step(scenario_name, TestID, Test_Name, Test_Description, Step_Name, Step_Description, Expected_Results, Role, Workstream)
        test.add_step(step)

    return tests


def create_word_documents_standalone(tests, folder_path="C:/Users/olive/Desktop/UAT_Test_Evidence", baseName=""):
    os.makedirs(folder_path, exist_ok=True)

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'
    doc.styles['Normal'].font.size = Pt(12)

    # Set the document to landscape orientation
    section = doc.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Add a Table of Contents
    p = doc.add_paragraph()
    run = p.add_run()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    run._r.append(fldSimple)

    # Test instructions
    p = doc.add_paragraph()
    run = p.add_run("Test Instructions\nPlease click ")
    run.bold = True
    run.font.size = Pt(14)

    # Adding the hyperlink
    hyperlink_url = 'https://myiglo.sharepoint.com/:w:/s/PhoenixProgramme2021-2024/EQ7d0Uhfs85Pit2Gk8DDnKsBP4Z25YQGJS34ULGxc39yew?e=KQHY5u'
    hyperlink = doc.part.relate_to(hyperlink_url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the hyperlink element
    hyperlink_element = OxmlElement('w:hyperlink')
    hyperlink_element.set(qn('r:id'), hyperlink)

    # Create the run for the hyperlink text
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set underline and color to blue
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = "UAT Test Instructions.docx"
    r.append(t)
    hyperlink_element.append(r)

    p._element.append(hyperlink_element)

    # Continue the sentence after the hyperlink
    run = p.add_run(" and read the instructions before you start testing!")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("UAT Test Script Execution")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Initialize table with a single row for headers
    table1 = doc.add_table(rows=1, cols=5)
    table1.style = 'Table Grid'

    # Apply header style
    for cell in table1.rows[0].cells:
        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="0047AB"/>'.format(nsdecls('w'))))

    # Define header cells text
    hdr_cells = table1.rows[0].cells
    hdr_text = ['Business Area', 'Responsible Tester', 'Status', 'Date', 'Country']
    for i, text in enumerate(hdr_text):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Define areas
    areas = ['Demand Planning', 'Supply Planning', 'Procurement', 'Warehousing', 'Quality Management', 'Production Planning', 'Plant maintenance', 'Commercial Finance', 'Order to Cash', 'FP&A-S4', 'R2R', 'MDG']

    # Add rows for each area and fill the first column with area names
    for area in areas:
        # Add a new row at the end of the table
        row = table1.add_row()
        # Set the text for the first cell in the row to the area name
        row.cells[0].text = area
            
    i = 0
    imax = len(tests)

    for test in tests:
        i += 1
        print(f"{i}/{imax}")
    
        doc.add_page_break()
    
        # Loop through the list of tests in the scenario
        if test.steps:
            first_step = test.steps[0]
            doc.add_heading(f"{first_step.TestName}", level=1)
            p = doc.add_paragraph()
            run = p.add_run(f"Test ID: {first_step.TestID} - {first_step.Workstream}\n\n")
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run("Test Description:\n")
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(f"{first_step.TestDescription}\n\n")
            run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            StepNumber = 0
            for step in test.steps:
                StepNumber += 1
                # Create the table with 5 rows and 6 columns
                table2 = doc.add_table(rows=5, cols=6)
                table2.style = 'Table Grid'

                # First row (Header Row)
                hdr_cells = table2.rows[0].cells

                # Apply bold formatting to specific headers
                hdr_cells[0].text = 'Test Step'
                hdr_cells[0].paragraphs[0].runs[0].bold = True

                hdr_cells[1].text = str(StepNumber)  # The test step number

                hdr_cells[2].text = 'Role'
                hdr_cells[2].paragraphs[0].runs[0].bold = True

                hdr_cells[3].text = step.Role  # Empty cell for Role data

                hdr_cells[4].text = 'Test Status <Pass/Fail>'
                hdr_cells[4].paragraphs[0].runs[0].bold = True

                hdr_cells[5].text = ''  # Empty cell for Test Status data

                # Second row
                row2_cells = table2.rows[1].cells
                row2_cells[0].text = 'Description'
                row2_cells[0].paragraphs[0].runs[0].bold = True
                row2_cells[1].merge(row2_cells[5])  # Merge the remaining cells to form a single cell
                row2_cells[1].text = step.StepDescription

                # Third row
                row3_cells = table2.rows[2].cells
                row3_cells[0].text = 'Expected'
                row3_cells[0].paragraphs[0].runs[0].bold = True
                row3_cells[1].merge(row3_cells[5])  # Merge the remaining cells to form a single cell
                row3_cells[1].text = step.ExpectedResults

                # Fourth row
                row3_cells = table2.rows[3].cells
                row3_cells[0].text = 'Actual Result'
                row3_cells[0].paragraphs[0].runs[0].bold = True
                row3_cells[1].merge(row3_cells[5])  # Merge the remaining cells to form a single cell
                row3_cells[1].text = ''

                # Fith row
                row4_cells = table2.rows[4].cells
                row4_cells[0].text = 'Tester Comments'
                row4_cells[0].paragraphs[0].runs[0].bold = True
                row4_cells[1].merge(row4_cells[5])  # Merge the remaining cells to form a single cell
                row4_cells[1].text = ''

                # Adjust alignment for all cells
                for row in table2.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
           
                doc.add_paragraph("\nPlease attach the test Screenshot Below, make sure to include the system date and time:\n\n\n\n\n", style='Normal')
        
    # Save document
    file_name_safe = f"{baseName[:100]}.docx"  # Truncate to ensure the filename is safe
    file_path = os.path.join(folder_path, file_name_safe)
    doc.save(file_path)

def process_all_files(excel_files_directory, word_files_directory):
    for filename in os.listdir(excel_files_directory):
        file_path = os.path.join(excel_files_directory, filename)
        if os.path.isfile(file_path) and filename.endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
            base_name = os.path.splitext(filename)[0]
            create_word_documents_standalone(read_excel_to_tests(file_path), word_files_directory, base_name)

# Get the current directory for Excel files
current_directory = os.getcwd()

# Get the parent directory for Word files
parent_directory = os.path.abspath(os.path.join(current_directory, os.pardir))

process_all_files(current_directory, parent_directory)
print("\n\n\t\tComplete")
